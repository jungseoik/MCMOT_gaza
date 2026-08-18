#!/usr/bin/env python3
"""사용 가이드 원고(마크다운) → Word(.docx) → PDF 빌드.

원고를 단일 진실 원천으로 두고 docx/pdf는 산출물로 만든다. docx는 바이너리라
변경 추적이 안 되므로, 문서 수정은 항상 `원고.md`에서 하고 이 도구로 다시 빌드한다.

지원 문법 (원고.md):
  @meta ... @end        표지 정보 (version/date/subtitle/org)
  @revision ... @end    개정 이력 표 (헤더 행 포함)
  # 제목                장 (페이지 나눔 + 목차 1수준)
  ## 제목               절 (목차 2수준)
  ### 제목              항
  | a | b |             표. 첫 행이 헤더
  @img 파일명 / 캡션 / @end    그림 + 번호 캡션 (그림 N-M)
  @note ... @end        주의 상자
  @steps ... @end       순서 목록
  @flow ... @end        흐름 표시(등폭)
  - 항목                글머리 목록
  일반 문단

사용:
  conda run -n boosttrack python tools/build_manual.py docs/manual/v0.0.1/원고.md
  conda run -n boosttrack python tools/build_manual.py <원고> --no-pdf

필요: python-docx (pip install python-docx) · PDF 변환은 libreoffice(soffice).
"""
from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "NanumGothic"
FONT_MONO = "D2Coding"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x60, 0x60, 0x60)
ACCENT = RGBColor(0x1C, 0x5C, 0xAB)
NOTE_BG = "F2F6FC"
HEAD_BG = "E8EDF5"
BODY_W_CM = 16.0                     # A4(21cm) - 좌우 여백 2.5cm씩


# ────────────────────────────────────────────────────── 원고 파싱
def parse(md: str) -> tuple[dict, list, list]:
    meta, revision, blocks = {}, [], []
    lines = md.split("\n")
    i, n = 0, len(lines)

    def gather(end="@end") -> list[str]:
        nonlocal i
        out = []
        i += 1
        while i < n and lines[i].strip() != end:
            out.append(lines[i])
            i += 1
        i += 1
        return out

    while i < n:
        ln = lines[i]
        s = ln.strip()

        if s == "@meta":
            for row in gather():
                if ":" in row:
                    k, _, v = row.partition(":")
                    meta[k.strip()] = v.strip()
            continue
        if s == "@revision":
            for row in gather():
                if row.strip().startswith("|"):
                    revision.append([c.strip() for c in
                                     row.strip().strip("|").split("|")])
            continue
        if s == "@note":
            blocks.append(("note", gather()))
            continue
        if s == "@steps":
            blocks.append(("steps", [x.strip() for x in gather() if x.strip()]))
            continue
        if s == "@flow":
            blocks.append(("flow", [x.strip() for x in gather() if x.strip()]))
            continue
        if s.startswith("@img"):
            src = s.split(None, 1)[1].strip()
            cap = gather()
            blocks.append(("img", (src, " ".join(x.strip() for x in cap).strip())))
            continue
        if s.startswith("### "):
            blocks.append(("h3", s[4:].strip())); i += 1; continue
        if s.startswith("## "):
            blocks.append(("h2", s[3:].strip())); i += 1; continue
        if s.startswith("# "):
            blocks.append(("h1", s[2:].strip())); i += 1; continue
        if s.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in
                             lines[i].strip().strip("|").split("|")])
                i += 1
            blocks.append(("table", rows))
            continue
        if s.startswith("- "):
            items = []
            while i < n and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(("ul", items))
            continue
        if s:
            para = [s]
            i += 1
            while i < n and lines[i].strip() and not re.match(
                    r"^\s*(#|\||-\s|@)", lines[i]):
                para.append(lines[i].strip()); i += 1
            blocks.append(("p", " ".join(para)))
            continue
        i += 1
    return meta, revision, blocks


# ────────────────────────────────────────────────────── docx 헬퍼
def shade(cell, hexcolor: str):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def set_font(run, size=10.5, bold=False, color=INK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def rich(par, text: str, size=10.5, color=INK):
    """**굵게** 와 `등폭` 만 처리하는 최소 인라인 서식."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_font(par.add_run(part[2:-2]), size, True, color)
        elif part.startswith("`") and part.endswith("`"):
            set_font(par.add_run(part[1:-1]), size - 0.5, False, ACCENT, FONT_MONO)
        else:
            set_font(par.add_run(part), size, False, color)


def field(par, code: str):
    """TOC 등 Word 필드 삽입."""
    r = par.add_run()
    for tag, attr, txt in (("w:fldChar", "w:fldCharType", "begin"),
                           ("w:instrText", "xml:space", "preserve"),
                           ("w:fldChar", "w:fldCharType", "separate"),
                           ("w:fldChar", "w:fldCharType", "end")):
        el = OxmlElement(tag)
        if tag == "w:instrText":
            el.set(qn(attr), "preserve")
            el.text = code
        else:
            el.set(qn(attr), txt)
        r._r.append(el)


def page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    set_font(r, 9, color=MUTED)
    for tag, attr, txt in (("w:fldChar", "w:fldCharType", "begin"),
                           ("w:instrText", "xml:space", "preserve"),
                           ("w:fldChar", "w:fldCharType", "end")):
        el = OxmlElement(tag)
        if tag == "w:instrText":
            el.set(qn(attr), "preserve"); el.text = "PAGE"
        else:
            el.set(qn(attr), txt)
        r._r.append(el)


def add_table(doc, rows: list[list[str]], header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    ncol = len(rows[0])
    first_w = Cm(BODY_W_CM * (0.28 if ncol > 2 else 0.34))
    rest = Cm((BODY_W_CM - first_w.cm) / max(1, ncol - 1))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.width = first_w if ci == 0 else rest
            par = cell.paragraphs[0]
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
            rich(par, val, size=9.5)
            if header and ri == 0:
                shade(cell, HEAD_BG)
                for r in par.runs:
                    r.bold = True
    return t


# ────────────────────────────────────────────────────── 빌드
def build(src: Path, out_docx: Path) -> dict:
    meta, revision, blocks = parse(src.read_text(encoding="utf-8"))
    img_dir = src.parent / "img"

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)   # A4 (기본값 Letter)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.5)

    title = blocks[0][1] if blocks and blocks[0][0] == "h1" else src.stem

    # ── 표지
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(meta.get("subtitle", "")), 11, color=MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), 24, True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(meta.get("version", "")), 13, True, ACCENT)
    for _ in range(10):
        doc.add_paragraph()
    for line in (meta.get("date", ""), meta.get("org", "")):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(line), 10.5, color=MUTED)

    # ── 개정 이력
    doc.add_page_break()
    p = doc.add_paragraph(); set_font(p.add_run("개정 이력"), 16, True)
    if revision:
        add_table(doc, revision)

    # ── 목차 (Word에서 F9로 갱신)
    doc.add_page_break()
    p = doc.add_paragraph(); set_font(p.add_run("목차"), 16, True)
    p = doc.add_paragraph()
    field(p, r'TOC \o "1-2" \h \z \u')
    p = doc.add_paragraph()
    set_font(p.add_run("※ 목차는 Word에서 F9(필드 업데이트)로 갱신됩니다."),
             9, color=MUTED)

    page_number_footer(sec)

    # ── 본문
    doc.add_page_break()
    chap, fig_no, stats = 0, 0, {"img": 0, "table": 0, "note": 0}
    skip_first_h1 = True

    for kind, val in blocks:
        if kind == "h1":
            if skip_first_h1:
                skip_first_h1 = False
                continue
            chap += 1
            fig_no = 0
            if chap > 1:
                doc.add_page_break()
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(10)
            set_font(h.add_run(val), 18, True, ACCENT)
        elif kind == "h2":
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            set_font(h.add_run(val), 13.5, True)
        elif kind == "h3":
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            set_font(h.add_run(val), 11.5, True)
        elif kind == "p":
            rich(doc.add_paragraph(), val)
        elif kind == "ul":
            for it in val:
                par = doc.add_paragraph(style="List Bullet")
                par.paragraph_format.space_after = Pt(2)
                rich(par, it)
        elif kind == "steps":
            for it in val:
                par = doc.add_paragraph()
                par.paragraph_format.left_indent = Cm(0.6)
                par.paragraph_format.space_after = Pt(3)
                rich(par, it)
        elif kind == "flow":
            for it in val:
                par = doc.add_paragraph()
                par.paragraph_format.left_indent = Cm(0.6)
                par.paragraph_format.space_after = Pt(2)
                set_font(par.add_run(it), 10, False, ACCENT, FONT_MONO)
        elif kind == "table":
            add_table(doc, val)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            stats["table"] += 1
        elif kind == "note":
            t = doc.add_table(rows=1, cols=1)
            t.style = "Table Grid"
            cell = t.rows[0].cells[0]
            cell.width = Cm(BODY_W_CM)
            shade(cell, NOTE_BG)
            first = True
            for ln in val:
                if not ln.strip():
                    continue
                par = cell.paragraphs[0] if first else cell.add_paragraph()
                par.paragraph_format.space_after = Pt(3)
                if first:
                    set_font(par.add_run("참고  "), 10, True, ACCENT)
                    first = False
                rich(par, ln.strip(), size=10)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            stats["note"] += 1
        elif kind == "img":
            src_name, caption = val
            path = img_dir / src_name
            if not path.is_file():
                print(f"  ⚠ 그림 없음: {src_name}")
                continue
            fig_no += 1
            stats["img"] += 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.add_run().add_picture(str(path), width=Cm(BODY_W_CM))
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            set_font(p.add_run(f"[그림 {chap}-{fig_no}] {caption}"), 9, color=MUTED)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    stats["chapters"] = chap
    return stats


def to_pdf(docx_path: Path) -> Path | None:
    r = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         str(docx_path), "--outdir", str(docx_path.parent)],
        capture_output=True, text=True, timeout=600)
    pdf = docx_path.with_suffix(".pdf")
    if not pdf.is_file():
        print("PDF 변환 실패:", r.stderr.strip()[:300], file=sys.stderr)
        return None
    return pdf


def main() -> int:
    ap = argparse.ArgumentParser(description="가이드 원고 → docx → pdf")
    ap.add_argument("src", type=Path, help="원고 마크다운 경로")
    ap.add_argument("--out", type=Path, help="출력 docx 경로 (기본: 원고와 같은 폴더)")
    ap.add_argument("--no-pdf", action="store_true")
    a = ap.parse_args()

    if not a.src.is_file():
        print(f"원고를 찾을 수 없습니다: {a.src}", file=sys.stderr)
        return 2

    meta, _, blocks = parse(a.src.read_text(encoding="utf-8"))
    title = blocks[0][1] if blocks and blocks[0][0] == "h1" else a.src.stem
    ver = meta.get("version", "v0.0.1")

    # 버전 폴더 = 스냅샷. 원고의 version과 폴더명이 어긋나면 개정 절차를 놓친 것이다
    # (개정은 폴더를 통째로 복사한 뒤 그 안에서 원고·화면을 갱신한다).
    if a.out is None and a.src.parent.name != ver:
        print(f"원고의 version({ver})과 폴더명({a.src.parent.name})이 다릅니다.\n"
              f"개정판은 폴더를 복사한 뒤 그 안에서 작업하십시오:\n"
              f"  cp -r {a.src.parent} {a.src.parent.parent / ver}\n"
              f"  python tools/capture_manual_shots.py "
              f"--out {a.src.parent.parent / ver}/img\n"
              f"  python tools/build_manual.py {a.src.parent.parent / ver}/원고.md",
              file=sys.stderr)
        return 2
    out = a.out or a.src.parent / f"{title.replace(' ', '-')}_{ver}.docx"

    print(f"▶ 빌드: {a.src}  →  {ver}")
    stats = build(a.src, out)
    print(f"  장 {stats['chapters']} · 그림 {stats['img']} · 표 {stats['table']}"
          f" · 참고 {stats['note']}")
    print(f"✅ {out}  ({out.stat().st_size/1024/1024:.1f} MB)")

    if not a.no_pdf:
        pdf = to_pdf(out)
        if pdf:
            print(f"✅ {pdf}  ({pdf.stat().st_size/1024/1024:.1f} MB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
